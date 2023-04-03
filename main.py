import sys

import docx
from music21 import converter, note, chord
from docx.enum.table import WD_TABLE_ALIGNMENT


def midi_to_jianpu(midi_file):
    midi = converter.parse(midi_file)
    jianpu = ""

    key_sig = midi.analyze('key')
    if key_sig:
        key_sig = f"1={key_sig.tonic} "

    time_sig = midi.getTimeSignatures()[0]
    if time_sig:
        time_sig = f"1/{time_sig.numerator} "

    hands = []
    for part in midi.parts:
        hand = []
        current_bar = ""
        bar_count = 0
        for measure in part.getElementsByClass('Measure'):
            for element in measure:
                if isinstance(element, note.Note):
                    current_bar += convert_note_to_jianpu(element)
                elif isinstance(element, chord.Chord):
                    current_bar += convert_chord_to_jianpu(element)
                current_bar += " "
            current_bar += "| "
            bar_count += 1
            if bar_count % 5 == 0:
                hand.append(current_bar.rstrip())
                current_bar = ""
        if current_bar:
            hand.append(current_bar.rstrip())
        hands.append(hand)

    max_length = max([len(hand) for hand in hands])
    for i in range(max_length):
        top_row = hands[0][i] if i < len(hands[0]) else ""
        bottom_row = hands[1][i] if i < len(hands[1]) else ""
        max_row_length = max(len(top_row), len(bottom_row))
        jianpu += top_row.ljust(max_row_length, ' ') + "\n" + bottom_row.ljust(max_row_length, ' ') + "\n\n"

    return jianpu, key_sig, time_sig


def convert_duration_to_jianpu(duration_obj):
    type_to_value = {
        'whole': ' - - -',
        'half': ' -',
        'quarter': '',
        'eighth': '\u0332',
        '16th': '\u0333'
    }
    duration_type = duration_obj.type
    if duration_type in type_to_value:
        base_duration = type_to_value[duration_type]

        if duration_obj.dots > 0:
            return base_duration + "•" * duration_obj.dots
        elif duration_type == 'half' and duration_obj.dots == 2:
            return ' - -\u0332'

        return base_duration

    return ''


def convert_note_to_jianpu(note_obj):
    scale_degrees = {
        0: '1',
        2: '2',
        4: '3',
        5: '4',
        7: '5',
        9: '6',
        11: '7'
    }

    octave_change = note_obj.octave - 4
    unicode_top = '\u0307'
    unicode_bottom = '\u0323'

    scale_degree = (note_obj.pitch.pitchClass - 0) % 12
    if scale_degree in scale_degrees:
        note_jianpu = scale_degrees[scale_degree]
    else:
        return ''  # Ignore notes that don't belong to the major scale

    if octave_change > 0:
        note_jianpu += unicode_top * octave_change
    elif octave_change < 0:
        note_jianpu += unicode_bottom * abs(octave_change)

    note_duration = convert_duration_to_jianpu(note_obj.duration)
    if note_duration:
        note_jianpu += note_duration

    return note_jianpu


def convert_chord_to_jianpu(chord_obj):
    jianpu_chord = ""
    for note_obj in chord_obj:
        jianpu_chord += convert_note_to_jianpu(note_obj)
    return jianpu_chord


def create_docx_file(jianpu_notation, output_file):
    doc = docx.Document()
    doc.add_heading('Jianpu Notation', 0)
    jianpu_paragraph = doc.add_paragraph(jianpu_notation)
    jianpu_paragraph.style.font.name = 'Doulos SIL Cipher'
    jianpu_paragraph.style.font.size = docx.shared.Pt(12)
    doc.save(output_file)


def save_to_docx(jianpu_notation, title, output_filename):
    doc = docx.Document()
    doc.add_heading(title, level=1)

    hands = jianpu_notation.split("\n\n")
    hands = [hand.split("\n") for hand in hands]

    for i in range(0, len(hands), 2):
        top_row = hands[i]
        bottom_row = hands[i + 1] if i + 1 < len(hands) else []

        max_rows = max(len(top_row), len(bottom_row))
        for j in range(max_rows):
            table = doc.add_table(rows=2, cols=len(top_row))
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            table.style = 'Table Grid'

            for col_idx, bar in enumerate(top_row):
                table.cell(0, col_idx).text = bar
                table.cell(0, col_idx).width = docx.shared.Inches(1)

            for col_idx, bar in enumerate(bottom_row):
                table.cell(1, col_idx).text = bar
                table.cell(1, col_idx).width = docx.shared.Inches(1)

            for row in table.rows:
                for cell in row.cells:
                    for direction in ['top', 'bottom', 'left', 'right']:
                        border_elements = cell._element.xpath(f'.//w:{direction}')
                        if border_elements:
                            border = border_elements[0]
                            border.set(docx.oxml.ns.qn('w:val'), 'nil')
            if j < max_rows - 1:
                doc.add_paragraph('')

    doc.save(output_filename)


def save_to_html(jianpu_notation, title, key_sig, time_sig, output_filename):
    hands = jianpu_notation.split("\n\n")
    hands = [hand.split("\n") for hand in hands]

    with open(output_filename, "w", encoding="utf-8") as output_file:
        output_file.write("<!DOCTYPE html>\n")
        output_file.write("<html>\n<head>\n<meta charset=\"UTF-8\">\n")
        output_file.write("<title>{}</title>\n".format(title))
        output_file.write("<link rel=\"stylesheet\" href=\"./styles.css\">\n")
        output_file.write("</head>\n<body>\n")
        output_file.write("<h1>{}</h1>\n".format(title))
        output_file.write("<h2>{} {}</h2>\n".format(key_sig, time_sig))
        for i in range(0, len(hands), 2):
            top_row = hands[i]
            bottom_row = hands[i + 1] if i + 1 < len(hands) else []

            max_rows = max(len(top_row), len(bottom_row))

            for j in range(0, max_rows, 5):
                output_file.write("<table>\n")

                # Top hand
                for k in range(j, min(j + 5, len(top_row))):
                    if top_row[k]:
                        output_file.write("<tr class=\"top-hand\">\n")
                        top_bar = top_row[k]
                        for top_note in top_bar.split("|")[:-1]:
                            output_file.write("<td class=\"bar\">{}</td>\n".format(top_note.strip().center(8)))
                        output_file.write("</tr>\n")

                output_file.write(
                    "<tr><td colspan=\"5\" style=\"border-top: none; border-bottom: none;\"></td></tr>\n")  # No-border between hands

                # Bottom hand
                for k in range(j, min(j + 5, len(bottom_row))):
                    if bottom_row[k]:
                        output_file.write("<tr class=\"bottom-hand\">\n")
                        bottom_bar = bottom_row[k]
                        for bottom_note in bottom_bar.split("|")[:-1]:
                            output_file.write("<td class=\"bar\">{}</td>\n".format(bottom_note.strip().center(8)))
                        output_file.write("</tr>\n")

                output_file.write("</table>\n")
                output_file.write("<br>\n")  # Double-sized space between sets of 5 bars

        output_file.write("</body>\n</html>")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Please provide the MIDI file path as an argument.")
        sys.exit(1)

    midi_file = sys.argv[1]
    jianpu_notation, key_sig, time_sig = midi_to_jianpu(midi_file)
    print(jianpu_notation)

    output_file = "jianpu_notation.docx"
    title = sys.argv[1]
    # save_to_docx(jianpu_notation, title, output_file)
    # save_to_text(jianpu_notation, "text.txt")
    save_to_html(jianpu_notation, title, key_sig, time_sig, "text.html")
    print(f"Jianpu notation has been saved to {output_file}")
